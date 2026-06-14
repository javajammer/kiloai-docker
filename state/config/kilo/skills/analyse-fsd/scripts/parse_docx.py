#!/usr/bin/env python3
"""
Parse .docx files and extract structured text content.

Output: JSON with metadata, sections, tables, hyperlinks, and full text.
Usage: python parse_docx.py <filepath> [--compact]
"""
import sys
import json
import os
from pathlib import Path


def extract_metadata(doc) -> dict:
    """Extract document metadata from core properties."""
    props = doc.core_properties
    return {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision or 0,
        "category": props.category or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "description": props.comments or "",
    }


def extract_hyperlinks(doc) -> list:
    """Extract all hyperlinks from the document."""
    hyperlinks = []
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                hyperlinks.append({
                    "target": rel.target_ref,
                    "text": rel.target_ref,
                })
    except Exception:
        pass
    return hyperlinks


def parse_docx(filepath: str, compact: bool = False) -> dict:
    """Parse a .docx file and return structured content."""
    try:
        from docx import Document
    except ImportError:
        return {
            "error": "python-docx not installed. Run: bash scripts/install_deps.sh",
            "format": "docx",
            "filepath": filepath,
        }

    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}", "format": "docx"}

    try:
        doc = Document(filepath)
    except Exception as e:
        return {"error": f"Failed to parse DOCX: {str(e)}", "format": "docx"}

    # Extract metadata
    metadata = extract_metadata(doc)

    # Extract sections with heading hierarchy
    sections = []
    current_section = {"heading": "Preamble", "level": 0, "content": [], "paragraphs": []}
    full_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text.append(text)

        if para.style.name.startswith("Heading"):
            # Save previous section
            if current_section["content"] or current_section["paragraphs"]:
                sections.append(current_section)

            # Parse heading level
            try:
                level = int(para.style.name.replace("Heading", "").strip())
            except ValueError:
                level = 1

            current_section = {
                "heading": text,
                "level": level,
                "content": [],
                "paragraphs": [],
            }
        else:
            current_section["content"].append(text)
            current_section["paragraphs"].append({
                "text": text,
                "style": para.style.name,
                "is_bold": any(run.bold for run in para.runs if run.bold),
                "is_italic": any(run.italic for run in para.runs if run.italic),
            })

    # Don't forget the last section
    if current_section["content"] or current_section["paragraphs"]:
        sections.append(current_section)

    # Extract tables with structured format
    tables = []
    for i, table in enumerate(doc.tables):
        table_data = {
            "index": i,
            "rows": [],
            "headers": [],
            "cell_count": 0,
        }

        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_data["rows"].append(cells)
            table_data["cell_count"] += len(cells)

            # First row is typically the header
            if j == 0:
                table_data["headers"] = cells

        tables.append(table_data)

    # Extract hyperlinks
    hyperlinks = extract_hyperlinks(doc)

    # Count statistics
    word_count = len(" ".join(full_text).split())
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    heading_count = len([s for s in sections if s["heading"] != "Preamble"])

    result = {
        "format": "docx",
        "filepath": filepath,
        "filename": os.path.basename(filepath),
        "metadata": metadata,
        "statistics": {
            "total_paragraphs": para_count,
            "total_headings": heading_count,
            "total_tables": len(tables),
            "total_sections": len(sections),
            "total_words": word_count,
            "total_characters": len(" ".join(full_text)),
        },
        "sections": sections,
        "tables": tables,
        "hyperlinks": hyperlinks,
        "full_text": "\n".join(full_text),
    }

    if compact:
        # Remove verbose fields for compact output
        for section in result["sections"]:
            section.pop("paragraphs", None)
        for table in result["tables"]:
            table.pop("cell_count", None)

    return result


def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_docx.py <filepath> [--compact]")
        print("")
        print("Parses .docx files and outputs structured JSON.")
        print("")
        print("Options:")
        print("  --compact    Remove verbose paragraph details from output")
        sys.exit(1)

    filepath = sys.argv[1]
    compact = "--compact" in sys.argv

    result = parse_docx(filepath, compact=compact)
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
